from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SRC = Path(r"C:\Users\linas\Downloads\1AtaskaitosTurinys (2).docx")
OUT = ROOT / "Keliones_ataskaita_uzpildyta.docx"
DIAGRAM_DIR = ROOT / "docs" / "generated"


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def add_after(anchor: Paragraph, text: str, style: str | None = None, bold: bool = False, color: tuple[int, int, int] | None = None) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, "", style)
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_picture_after(anchor: Paragraph, path: Path, width: float = 6.4) -> Paragraph:
    paragraph = insert_paragraph_after(anchor)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    paragraph.paragraph_format.space_after = Pt(8)
    return paragraph


def move_table_after(table: Table, anchor: Paragraph) -> None:
    element = table._tbl
    element.getparent().remove(element)
    anchor._p.addnext(element)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def in_sequence_section(doc: Document) -> list[Paragraph]:
    inside = False
    result: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Panaudojimo atvejų sekų diagramos":
            inside = True
            result.append(paragraph)
            continue
        if inside and text == "Sistemos klasių diagrama":
            break
        if inside:
            result.append(paragraph)
    return result


def fill_report() -> Path:
    shutil.copyfile(SRC, OUT)
    doc = Document(OUT)

    summaries = {
        "P1 Redaguoti maršrutą:": "Realizuota maršrutų puslapyje: naudotojas pasirenka esamą maršrutą, mato OSM žemėlapį, prideda arba pašalina POI taškus ir perskaičiuoja maršrutą per pasirinktus taškus.",
        "P2 Peržiūrėti maršrutus:": "Realizuota kaip išsaugotų maršrutų sąrašas su atnaujinimu, redagavimu ir šalinimu. Duomenys gaunami iš RouteController.getAllRoutes().",
        "P3 Sukurti maršrutą:": "Realizuota kaip naujo maršruto sudarymas iš pradžios miesto, pabaigos miesto ir pasirinktų POI tarp miestų. Maršrutas skaičiuojamas per RouteController.saveRoute().",
        "P4 Pasirinkti apgyvendinimą iš sąrašo:": "Realizuota kelionių puslapyje: pasirinkus aktyvią kelionę rodoma Hotels skiltis, o pasirinktas viešbutis išsaugomas per assignAccommodationToTrip().",
        "P5 Pasirinkti skrydį iš sąrašo:": "Realizuota kelionių puslapyje: Flights skiltyje pateikiami skrydžiai pagal pasirinkto maršruto miestus, pasirinkimas išsaugomas kelionėje.",
        "P6 Pasirinkti automobilį iš sąrašo:": "Realizuota kelionių puslapyje: Cars skiltyje pateikiami automobilio pasiūlymai pagal kelionės tikslą, pasirinkimas susiejamas su aktyvia kelione.",
        "P7 Pasirinkti maršrutą:": "Realizuota kelionių puslapyje: kelionė nebekuria maršruto, o tik pasirenka jau išsaugotą maršrutą iš RouteController.getAllRoutes().",
        "P11 Rasti lankytinus objektus:": "Realizuota per RouteController.getRoadPOI(): gaunami OSM POI tarp miestų, o kai išoriniai šaltiniai nepasiekiami, sukuriami vietiniai koridoriaus POI.",
        "P12 Valdyti kelionės rezervacijas:": "Viešbučio, skrydžio ir automobilio pasirinkimai sukuria rezervacijos įrašus per createReservation(), todėl kelionė turi susietus rezervuojamus elementus.",
        "P13 Įvertinti kelionės trukmę:": "Realizuota RouteController.saveRouteTime(): trukmė apskaičiuojama iš OSRM arba vietinės atsarginės geometrijos ir išsaugoma maršrute.",
        "P14 Valdyti keliones:": "Realizuota kelionių puslapyje: sukuriama kelionė iš pasirinkto maršruto, datų ir paslaugų; maršruto redagavimas atskirtas į maršrutų puslapį.",
    }

    section_paragraphs = list(in_sequence_section(doc))
    for paragraph in reversed(section_paragraphs):
        text = paragraph.text.strip()
        if text in summaries:
            inserted = insert_paragraph_after(paragraph, summaries[text], "Normal")
            inserted.paragraph_format.space_after = Pt(6)

    # Insert required diagrams and mapping after M3 slot.
    section_paragraphs = list(in_sequence_section(doc))
    anchor = next((p for p in section_paragraphs if p.text.strip() == "M3 Peržiūrėti mokėjimų istoriją:"), section_paragraphs[-1])

    anchor = add_after(anchor, "Papildytos sekų diagramos pagal realizuotą prototipą", "Heading 3")
    anchor = add_after(
        anchor,
        "Šiame poskyryje pateikiama, kaip turi atrodyti realizuotų dalių sekų diagramos. "
        "Maršrutų ir kelionių atsakomybės atskirtos: maršrutai apima OSM žemėlapį, POI pasirinkimą ir perskaičiavimą, "
        "o kelionės apima tik maršruto, viešbučio, skrydžio ir automobilio pasirinkimą.",
        "Normal",
    )

    diagrams = [
        ("P1/P3/P11/P13 Maršruto redagavimas ir perskaičiavimas", "route_editor_sequence.png"),
        ("P4/P7/P14 Kelionės sukūrimas iš pasirinkto maršruto", "trip_selection_sequence.png"),
        ("P4/P5/P6/P12 Viešbučio, skrydžio ir automobilio pasirinkimas", "service_selection_sequence.png"),
    ]
    for title, image_name in diagrams:
        anchor = add_after(anchor, title, "Normal", bold=True, color=(37, 99, 235))
        anchor = add_picture_after(anchor, DIAGRAM_DIR / image_name)

    anchor = add_after(anchor, "Diagramos ir kodo atitikimo lentelė", "Heading 3")
    anchor = add_after(anchor, "Lentelėje nurodoma, kuri realizuoto kodo dalis atitinka sekų diagramų pranešimus ir valdiklius.", "Normal")

    rows = [
        ("P1/P3 Sukurti arba redaguoti maršrutą", "RoutesPage.tsx, RouteMap.tsx", "POST/PUT /api/Route", "saveRoute, sendCities, saveRouteData"),
        ("P11 Rasti lankytinus objektus", "RoutesPage.findObjects", "POST /api/Route/roadPOI", "getRoadPOI, getOSMData, createCorridorPOI"),
        ("POI pasirinkimas žemėlapyje", "RouteMap marker click", "routePoints payload", "selectObjects, createRouteWithPOI"),
        ("Maršruto perskaičiavimas per POI", "RoutesPage.saveRoute", "RouteController.saveRoute/sendRouteData", "createLengthMatrix, shuffleObjectOrder, createPolyLine"),
        ("P2 Peržiūrėti maršrutus", "RoutesPage.getRoutes", "GET /api/Route", "getAllRoutes, getSpecificRoute"),
        ("P7 Pasirinkti maršrutą", "TripsPage route selector", "GET /api/Route", "getAllRoutes"),
        ("P14 Valdyti keliones", "TripsPage.createTripFromSelection", "POST /api/Trip", "saveTripInformation, Trip.checkTrip"),
        ("P4 Pasirinkti apgyvendinimą", "TripsPage Hotels tab", "POST /api/Trip/accommodation/list; POST /api/Trip/{id}/accommodation", "requestAccommodationListFromExternalActor, assignAccommodationToTrip"),
        ("P5 Pasirinkti skrydį", "TripsPage Flights tab", "POST /api/Trip/flight/list; POST /api/Trip/{id}/flight", "requestFlightListFromExternalActor, assignFlightToTrip"),
        ("P6 Pasirinkti automobilį", "TripsPage Cars tab", "POST /api/Trip/car/list; POST /api/Trip/{id}/car", "requestCarListFromExternalActor, assignCarToTrip"),
        ("P12 Valdyti rezervacijas", "OfferColumn select buttons", "POST /api/Trip/{id}/...", "createReservation, saveSelectedAccommodation/Flight/Car"),
        ("P13 Įvertinti kelionės trukmę", "RouteMap statistics", "POST/PUT /api/Route; GET /api/Route/getTime/{id}", "saveRouteTime, evaluateRouteData"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Diagrama / dalis", "Frontend", "Endpoint / controller", "Funkcijos"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "D9E8FF")
        set_cell_text(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell_text(cell, value)
    move_table_after(table, anchor)

    proto_anchor = next((p for p in doc.paragraphs if p.text.strip() == "Sistemos prototipas"), doc.paragraphs[-1])
    proto_anchor = add_after(
        proto_anchor,
        "Realizuotas prototipas turi atskirus maršrutų ir kelionių puslapius. Maršrutų puslapyje naudojama Leaflet / OpenStreetMap biblioteka, todėl miestai ir POI taškai rodomi realiame OSM žemėlapyje. Kelionių puslapyje naudotojas pasirenka jau paruoštą maršrutą ir kelionės paslaugas: viešbutį, skrydį ir automobilį.",
        "Normal",
    )
    add_after(
        proto_anchor,
        "Pagrindiniai realizuoti failai: frontend/src/Views/Travel/RoutesPage.tsx, frontend/src/Views/Travel/RouteMap.tsx, frontend/src/Views/Travel/TripsPage.tsx, backend/Controllers/RouteController.cs, backend/Controllers/TripController.cs.",
        "Normal",
    )

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(fill_report())
